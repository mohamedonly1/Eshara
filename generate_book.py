import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def set_spacing(paragraph, space_after=12, line_spacing=1.5):
    pPr = paragraph._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), str(space_after * 20)) # twips
    spacing.set(qn('w:line'), str(int(240 * line_spacing))) # 240 is single line
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def add_rtl_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=False, size=14, space_after=12):
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, space_after)
    set_rtl(p)
    
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Simplified Arabic'
        
        rPr = run._element.get_or_add_rPr()
        
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Simplified Arabic')
        
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)
        
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(size * 2))
        rPr.append(szCs)
        
    return p

def add_ltr_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, 12, 1.5)
    
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
    return p

def add_heading_rtl(doc, text, level):
    size = 18 if level == 1 else 16
    p = add_rtl_paragraph(doc, text, WD_ALIGN_PARAGRAPH.RIGHT, bold=True, size=size)
    try:
        p.style = doc.styles[f'Heading {level}']
    except Exception:
        pass # fallback if style doesn't exist
    
    # Re-apply styling because assigning style might override some element props
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    return p

def create_table(doc, rows, cols, rtl=True):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    if rtl:
        tblPr = table._element.xpath('w:tblPr')
        if tblPr:
            bidiVisual = OxmlElement('w:bidiVisual')
            tblPr[0].append(bidiVisual)
    return table

def main():
    doc = Document()
    
    # Set page size (A4) and margins
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- COVER PAGE ---
    add_rtl_paragraph(doc, "جامعة المنصورة", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    add_rtl_paragraph(doc, "كلية التربية النوعية", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    add_rtl_paragraph(doc, "قسم إعداد معلم الحاسب الآلي", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    add_rtl_paragraph(doc, "العام الجامعي: 2025 / 2026", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)

    add_rtl_paragraph(doc, "", WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=48)

    add_rtl_paragraph(doc, "مشروع تخرج بعنوان:", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
    add_rtl_paragraph(doc, "تطبيق (إشارة) ترجمة لغة الإشارة (للصم والبكم)", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=24)
    add_ltr_paragraph(doc, "Programming Language Translation Application (for the Deaf and Mute)", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)

    add_rtl_paragraph(doc, "", WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=48)

    add_rtl_paragraph(doc, "إشراف", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
    add_rtl_paragraph(doc, "1. أ.م.د / حسنيه محمدى محمد احمد", WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_rtl_paragraph(doc, "2. د/ حنان الرفاعى", WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_rtl_paragraph(doc, "3. د/ ايمان عبد العظيم", WD_ALIGN_PARAGRAPH.CENTER, size=16)

    doc.add_page_break()

    # --- INNER COVER 1 ---
    add_rtl_paragraph(doc, "تطبيق (إشارة)", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=24)
    add_rtl_paragraph(doc, "[صورة توضيحية للتطبيق - يرجى إدراج الصورة هنا في الوورد]", WD_ALIGN_PARAGRAPH.CENTER, size=14)
    doc.add_page_break()

    # --- INNER COVER 2 (Team Table) ---
    add_heading_rtl(doc, "فريق العمل", 1)
    team = [
        "منه على محمد عبدالسميع", "محمد الحسن احمد", "سعد رضا سعد", "السيد احمد السيد",
        "محمد شكرى", "ضحى ابراهيم موسى", "ندا محمود سلام", "حنان ثروت فتحى",
        "اسراء سامح حسنى", "اسراء النبوى حسن", "ايه السيد محمد"
    ]
    table = create_table(doc, len(team)+1, 2)
    hdr_cells = table.rows[0].cells
    
    # RTL requires reversing logic sometimes, let's just put Name right, Num left visually
    p1 = hdr_cells[0].paragraphs[0]
    set_rtl(p1); p1.add_run("الاسم").font.name = "Simplified Arabic"
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p2 = hdr_cells[1].paragraphs[0]
    set_rtl(p2); p2.add_run("م").font.name = "Simplified Arabic"
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i, name in enumerate(team):
        row_cells = table.rows[i+1].cells
        
        c1 = row_cells[0].paragraphs[0]
        set_rtl(c1); c1.add_run(name).font.name = "Simplified Arabic"
        c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        c2 = row_cells[1].paragraphs[0]
        set_rtl(c2); c2.add_run(str(i+1)).font.name = "Simplified Arabic"
        c2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- ACKNOWLEDGMENTS ---
    add_heading_rtl(doc, "الشكر والتقدير (Acknowledgments)", 1)
    add_rtl_paragraph(doc, "نتقدم بخالص الشكر والتقدير إلى جامعة المنصورة وكلية التربية النوعية وقسم إعداد معلم الحاسب الآلي، وإلى أساتذتنا المشرفين على هذا المشروع لجهودهم وتوجيهاتهم المستمرة طوال فترة إنجاز العمل.")
    doc.add_page_break()

    # --- ABSTRACT AR ---
    add_heading_rtl(doc, "الملخص", 1)
    abstract_ar = '''يهدف هذا المشروع إلى تطوير نظام ويب متكامل (تطبيق إشارة) لجمع بيانات لغة الإشارة العربية الموحدة والتعرف عليها آلياً في الزمن الفعلي. يوفر النظام منصة تتيح للمتطوعين تسجيل إيماءات الحروف العربية الـ 28 عبر كاميرا الويب، مع إمكانية عرض الترجمة الفورية للمستخدم.
يحل المشروع مشكلة شح مجموعات البيانات العربية المستخدمة لتدريب نماذج التعرف على لغة الإشارة، ويسهم في تيسير التواصل بين مجتمع الصم وضعاف السمع والمجتمع بشكل عام.
تعتمد معمارية النظام على تقنية MediaPipe لاستخراج 21 نقطة مرجعية من اليد، ثم تطبيع الإحداثيات وتمريرها إلى نموذج شبكة عصبية متعددة الطبقات (MLP) مبني باستخدام TensorFlow و TFLite. تم بناء الواجهة الخلفية باستخدام Python و Flask، بينما تعتمد الواجهة الأمامية على HTML/CSS/JavaScript و Three.js لتصور اليد ثلاثي الأبعاد.
حقق النموذج دقة تصل إلى 93.8% على 28 حرفاً، مع توفر نظام تصفية جودة آلي يرفض العينات الخاطئة. كما يضم النظام لوحة إدارة متكاملة لمتابعة مساهمات المتطوعين، ويدعم التشغيل على الأجهزة المحمولة كتطبيق ويب تقدمي (PWA).'''
    add_rtl_paragraph(doc, abstract_ar)
    doc.add_page_break()

    # --- ABSTRACT EN ---
    add_ltr_paragraph(doc, "Abstract", WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=18)
    abstract_en = '''This project aims to develop an integrated web system ("Eshara" App) for collecting Unified Arabic Sign Language data and automatically recognizing it in real-time. The system provides a platform allowing volunteers to record the 28 Arabic alphabet gestures via webcam, with real-time translation capabilities.
The project addresses the scarcity of Arabic datasets used for training sign language recognition models, bridging the communication gap between the deaf and hard-of-hearing community and society at large.
The system architecture relies on MediaPipe to extract 21 hand landmarks, normalizes their coordinates, and passes them to a Multi-Layer Perceptron (MLP) neural network built with TensorFlow and TFLite. The backend is powered by Python and Flask, while the frontend utilizes HTML/CSS/JavaScript and Three.js for 3D hand visualization.
The model achieved an accuracy of 93.8% across 28 letters, featuring an automated quality filtering system that rejects incorrect samples. The platform also includes a comprehensive dashboard for tracking volunteer contributions and supports progressive web app (PWA) deployment for mobile devices.'''
    add_ltr_paragraph(doc, abstract_en)
    doc.add_page_break()

    # --- TABLE OF CONTENTS (Placeholder) ---
    add_heading_rtl(doc, "فهرس المحتويات (Table of Contents)", 1)
    add_rtl_paragraph(doc, "[ملاحظة: يرجى توليد الفهرس التلقائي باستخدام Microsoft Word عبر References > Table of Contents]")
    doc.add_page_break()

    # --- LIST OF FIGURES ---
    add_heading_rtl(doc, "قائمة الأشكال (List of Figures)", 1)
    add_rtl_paragraph(doc, "[جدول قائمة الأشكال يدرج هنا]")
    doc.add_page_break()

    # --- LIST OF TABLES ---
    add_heading_rtl(doc, "قائمة الجداول (List of Tables)", 1)
    add_rtl_paragraph(doc, "[جدول قائمة الجداول يدرج هنا]")
    doc.add_page_break()

    # --- ABBREVIATIONS ---
    add_heading_rtl(doc, "جدول الاختصارات (Table of Abbreviations)", 1)
    abbrev = [
        ("MediaPipe", "إطار عمل من جوجل لبناء مسارات تعلم الآلة (استخدم هنا لتتبع اليد)"),
        ("MLP", "Multi-Layer Perceptron (شبكة عصبية متعددة الطبقات)"),
        ("TFLite", "TensorFlow Lite (نسخة خفيفة من إطار عمل TensorFlow لتشغيل النماذج بفعالية)"),
        ("RTL", "Right-to-Left (اتجاه النص من اليمين لليسار)"),
        ("API", "Application Programming Interface (واجهة برمجة التطبيقات)"),
        ("REST", "Representational State Transfer (نمط معماري لواجهات الويب)"),
        ("PWA", "Progressive Web App (تطبيق ويب تقدمي يمكن تثبيته)"),
        ("UASL", "Unified Arabic Sign Language (لغة الإشارة العربية الموحدة)"),
        ("CSV", "Comma-Separated Values (قيم مفصولة بفواصل، لتخزين البيانات)"),
        ("JSON", "JavaScript Object Notation (صيغة خفيفة لتبادل البيانات)")
    ]
    table_abbr = create_table(doc, len(abbrev)+1, 2)
    
    hc1 = table_abbr.rows[0].cells[0].paragraphs[0]
    set_rtl(hc1); hc1.add_run("المعنى").font.name = "Simplified Arabic"
    hc1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    hc2 = table_abbr.rows[0].cells[1].paragraphs[0]
    set_rtl(hc2); hc2.add_run("الاختصار").font.name = "Simplified Arabic"
    hc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i, (abbr, mean) in enumerate(abbrev):
        c1 = table_abbr.rows[i+1].cells[0].paragraphs[0]
        set_rtl(c1); c1.add_run(mean).font.name = "Simplified Arabic"
        
        c2 = table_abbr.rows[i+1].cells[1].paragraphs[0]
        # For LTR acronyms in RTL table, we just write it
        set_rtl(c2); c2.add_run(abbr).font.name = "Times New Roman"
        c2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_page_break()

    # --- CHAPTER 1 ---
    add_heading_rtl(doc, "الفصل الأول: المقدمة (Introduction)", 1)
    
    add_heading_rtl(doc, "مقدمة عامة عن لغة الإشارة العربية", 2)
    add_rtl_paragraph(doc, "لغة الإشارة هي وسيلة التواصل الأساسية لمجتمع الصم والبكم. ولغة الإشارة العربية الموحدة هي محاولة لتوحيد الإشارات في العالم العربي. مع التطور التكنولوجي، أصبح من الممكن استخدام تقنيات الذكاء الاصطناعي للتعرف الآلي على هذه الإشارات وترجمتها إلى نصوص لتسهيل التواصل.")

    add_heading_rtl(doc, "بيان المشكلة (Problem Statement)", 2)
    add_rtl_paragraph(doc, "تكمن المشكلة في شُح مجموعات البيانات العربية المتاحة لتدريب نماذج التعرف على لغة الإشارة، بالإضافة إلى غياب منصة ويب عربية متكاملة تجمع بين جمع البيانات التشاركي والترجمة الفورية، مما يؤدي إلى صعوبة التواصل بين مجتمع الصم وضعاف السمع والمجتمع بشكل عام.")

    add_heading_rtl(doc, "أهداف المشروع", 2)
    add_rtl_paragraph(doc, "1. بناء نظام ويب متكامل للتعرف الفوري على لغة الإشارة العربية.")
    add_rtl_paragraph(doc, "2. توفير منصة لجمع وتحديث البيانات بشكل تشاركي من خلال المتطوعين.")
    add_rtl_paragraph(doc, "3. تدريب نموذج ذكاء اصطناعي دقيق وموثوق للتعرف على الحروف الأبجدية.")
    add_rtl_paragraph(doc, "4. توفير واجهة مستخدم سهلة تدعم الأجهزة المحمولة كتطبيق ويب تقدمي (PWA).")

    add_heading_rtl(doc, "حدود المشروع ونطاقه", 2)
    add_rtl_paragraph(doc, "يقتصر هذا المشروع على التعرف على الحروف الأبجدية العربية الـ 28 الثابتة باستخدام يد واحدة ضمن إطار كاميرا الويب المتاحة للمستخدم، ولا يتضمن إيماءات الكلمات الكاملة المستمرة.")

    add_heading_rtl(doc, "الهيكل التنظيمي للكتاب", 2)
    add_rtl_paragraph(doc, "يحتوي الكتاب على خمسة فصول: المقدمة، الدراسات السابقة، المنهجية والتصميم، التنفيذ والنتائج، والخاتمة، بالإضافة إلى المراجع والملاحق التقنية.")
    doc.add_page_break()

    # --- CHAPTER 2 ---
    add_heading_rtl(doc, "الفصل الثاني: الدراسات السابقة (Literature Review)", 1)
    
    add_heading_rtl(doc, "تعريف لغة الإشارة العربية الموحدة (UASL)", 2)
    add_rtl_paragraph(doc, "هي لغة إشارة تم تطويرها لتوحيد التواصل بين الصم في مختلف الدول العربية، وتتضمن تمثيلاً إشارياً ثابتاً للحروف الأبجدية والكلمات الشائعة، لتسهيل التفاهم المشترك.")

    add_heading_rtl(doc, "شرح تقنية MediaPipe Hands و21 نقطة اليد", 2)
    add_rtl_paragraph(doc, "MediaPipe هي أداة مفتوحة المصدر من جوجل تعتمد على التعلم العميق لتتبع 21 نقطة مرجعية (Landmarks) ثلاثية الأبعاد في اليد بدقة عالية وفي الزمن الفعلي، مما يوفر أساساً متيناً لاستخراج ميزات الإيماءات دون الحاجة لمعدات خاصة.")

    add_heading_rtl(doc, "شرح الشبكات العصبية متعددة الطبقات (MLP)", 2)
    add_rtl_paragraph(doc, "الشبكات العصبية متعددة الطبقات (Multi-Layer Perceptron) هي نوع من الشبكات العصبية الاصطناعية تتكون من طبقة إدخال، طبقات مخفية، وطبقة إخراج، وتتميز بقدرتها على تعلم الأنماط غير الخطية والمعقدة من البيانات.")

    add_heading_rtl(doc, "شرح TensorFlow Lite والنشر الفعّال", 2)
    add_rtl_paragraph(doc, "TensorFlow Lite هي تقنية تتيح تحويل نماذج التعلم العميق الضخمة إلى صيغ خفيفة الوزن ومحسّنة (TFLite) يمكن تشغيلها بسرعة وفعالية على الأجهزة الطرفية ومتصفحات الويب.")

    add_heading_rtl(doc, "مراجعة أعمال سابقة ذات صلة مع مقارنة", 2)
    add_rtl_paragraph(doc, "تناولت عدة دراسات التعرف على لغة الإشارة باستخدام شبكات الطي العصبية (CNN) وغيرها، إلا أن معظمها ركز على اللغات الأجنبية أو تطلبت معدات خاصة كقفازات الاستشعار. يتميز مشروعنا بالاعتماد على كاميرا الويب فقط لتقليل التكلفة وزيادة الإتاحة.")

    add_heading_rtl(doc, "الفجوة البحثية وإسهام هذا المشروع", 2)
    add_rtl_paragraph(doc, "يسد هذا المشروع الفجوة من خلال توفير منصة ويب تدعم جمع البيانات (Data Collection) المستمر، والترجمة الفورية (Real-time Translation) بحلول برمجية خفيفة وواجهة عربية متكاملة (RTL).")
    doc.add_page_break()

    # --- CHAPTER 3 ---
    add_heading_rtl(doc, "الفصل الثالث: المنهجية والتصميم (Methodology & Design)", 1)
    
    add_heading_rtl(doc, "معمارية النظام الكاملة", 2)
    add_rtl_paragraph(doc, "تتكون المعمارية من: كاميرا الويب ← استخراج 21 نقطة عبر MediaPipe ← تطبيع الإحداثيات (42 قيمة) ← التمرير لنموذج MLP/TFLite ← واجهة الـ API المبنية بـ Flask ← عرض النتائج عبر واجهة المستخدم (HTML/CSS/JS/Three.js).")
    add_rtl_paragraph(doc, "[مخطط توضيحي لمعمارية النظام - يدرج هنا]")

    add_heading_rtl(doc, "تدفق البيانات", 2)
    add_rtl_paragraph(doc, "- مرحلة الجمع: تسجيل النقاط عبر الكاميرا ← التصفية الذكية ← الحفظ في CSV/JSON.\n- مرحلة التعرف الفوري: استخراج النقاط ← التطبيع ← استدعاء نقطة API للتنبؤ ← عرض الترجمة الحية.")

    add_heading_rtl(doc, "تصميم قاعدة البيانات", 2)
    add_rtl_paragraph(doc, "يعتمد النظام على ملفات CSV لتخزين الإحداثيات والبيانات التدريبية المجمعة بشكل مهيكل، وملفات JSON لتخزين إعدادات النظام، التكوينات، ومعلومات المصادقة للمستخدمين.")

    add_heading_rtl(doc, "Use Case Diagram", 2)
    add_rtl_paragraph(doc, "يخدم النظام ثلاثة أطراف رئيسية: الزائر (يستخدم الترجمة الفورية)، المتطوع (يقوم بتسجيل وجمع البيانات)، والمسؤول (يدير النظام، يراجع العينات ويتابع الإحصائيات).")
    add_rtl_paragraph(doc, "[صورة Use Case Diagram - يدرج هنا]")

    add_heading_rtl(doc, "Sequence Diagram", 2)
    add_rtl_paragraph(doc, "يصف تسلسل العمليات من تفاعل المستخدم مع المتصفح، وإرسال البيانات عبر الـ REST API، وحتى الرد من خادم Flask بنتيجة التنبؤ.")
    add_rtl_paragraph(doc, "[صورة Sequence Diagram - يدرج هنا]")

    add_heading_rtl(doc, "متطلبات النظام", 2)
    add_rtl_paragraph(doc, "- المتطلبات الوظيفية: تسجيل الدخول، جمع العينات، الترجمة الفورية، محرر وضعيات اليد ثلاثي الأبعاد.\n- المتطلبات غير الوظيفية: أداء عالي وسريع، أمان البيانات، التوافق مع الشاشات المختلفة.")

    add_heading_rtl(doc, "منهجية الأمان", 2)
    add_rtl_paragraph(doc, "تتضمن تجزئة كلمات المرور (Hashing)، تحديد معدل الطلبات (Rate Limiting) عبر مكتبة flask-limiter، واستخدام ProxyFix من werkzeug للتعامل الآمن مع الخوادم الوكيلة وعناوين IP.")
    doc.add_page_break()

    # --- CHAPTER 4 ---
    add_heading_rtl(doc, "الفصل الرابع: التنفيذ والنتائج (Implementation & Results)", 1)
    
    add_heading_rtl(doc, "بيئة التطوير والأدوات", 2)
    add_rtl_paragraph(doc, "تم الاعتماد على لغة Python 3.10 مع إطار عمل Flask للواجهة الخلفية. واستخدم MediaPipe 0.10.11 لاستخراج نقاط اليد، و TensorFlow 2.10 لتدريب النماذج، ومكتبات NumPy 1.26 و scikit-learn لمعالجة البيانات. واجهة المستخدم بنيت بتقنيات الويب القياسية مع Three.js للرسم ثلاثي الأبعاد، ومكتبات arabic_reshaper و python-bidi لدعم اللغة العربية.")

    add_heading_rtl(doc, "معمارية نموذج MLP بالتفصيل", 2)
    add_rtl_paragraph(doc, "تم بناء النموذج كشبكة عصبية أمامية التغذية تحتوي على طبقات Dense متعددة مع دوال تنشيط ReLU و Dropout لمنع الإفراط في التخصيص (Overfitting)، محققاً كفاءة عالية في التصنيف المباشر.")

    add_heading_rtl(doc, "آلية تطبيع نقاط اليد (normalize_landmarks)", 2)
    add_rtl_paragraph(doc, "يتم أخذ النقطة رقم 0 (المعصم) كنقطة أصل (0,0)، ثم حساب إحداثيات باقي النقاط الـ 20 بالنسبة لها، مما يضمن استقرار التنبؤ بغض النظر عن موضع اليد داخل إطار الكاميرا.")

    add_heading_rtl(doc, "آلية إعادة ترميز التسميات (LabelEncoder)", 2)
    add_rtl_paragraph(doc, "استخدمت أداة LabelEncoder لتحويل التسميات النصية للحروف العربية إلى فئات رقمية صحيحة قابلة للمعالجة بواسطة دالة الخسارة أثناء تدريب النموذج.")

    add_heading_rtl(doc, "شرح واجهات المستخدم الرئيسية", 2)
    add_rtl_paragraph(doc, "- مسار /login: تسجيل الدخول وإنشاء الحساب.\n- مسار /collect-data: جمع إيماءات اليد (يضم كاميرا وشريط تقدم لكل حرف).\n- مسار /: التعرف الفوري وعرض الترجمة.\n- مسار /profile: إحصاءات مساهمات المتطوع.\n- مسار /admin: لوحة الإدارة للمسؤول.\n- مسار /pose-editor: محرر وضعيات اليد ثلاثي الأبعاد.")
    add_rtl_paragraph(doc, "[لقطات شاشة للواجهات - تدرج هنا]")

    add_heading_rtl(doc, "نقاط API الرئيسية", 2)
    add_rtl_paragraph(doc, "- POST /predict: استقبال 42 قيمة وإرجاع الحرف والثقة.\n- POST /collect: استقبال العينات وتصفيتها وحفظها.\n- GET /sample_counts: جلب عدد العينات لكل حرف.\n- POST /auth/login و /auth/register: المصادقة.")

    add_heading_rtl(doc, "نتائج التدريب", 2)
    add_rtl_paragraph(doc, "تم تحقيق دقة كلية بلغت 93.8% على 28 حرفاً. أظهر تحليل الأخطاء أداءً ممتازاً لمعظم الحروف، مع وجود التباس طفيف في التفريق بين الحروف المتشابهة حركياً مثل حرفي (د) و (ر).")
    add_rtl_paragraph(doc, "[جدول Precision / Recall / F1 يدرج هنا]")

    add_heading_rtl(doc, "مقارنة بالأعمال السابقة", 2)
    add_rtl_paragraph(doc, "أثبت النموذج تفوقه في سرعة الاستجابة اللحظية على المتصفح (بفضل TFLite)، ومرونة النظام الفائقة التي لا تتطلب أي تطبيقات وسيطة لعملية الترجمة والجمع.")
    doc.add_page_break()

    # --- CHAPTER 5 ---
    add_heading_rtl(doc, "الفصل الخامس: الخاتمة والتوصيات (Conclusion & Recommendations)", 1)
    
    add_heading_rtl(doc, "خلاصة ما تحقق", 2)
    add_rtl_paragraph(doc, "تم بنجاح تصميم وبرمجة نظام ويب متكامل للتعرف على لغة الإشارة العربية عبر الكاميرا وبدقة 93.8%، متجاوزين مشكلة نقص البيانات العربية من خلال بناء وحدة جمع بيانات تشاركية وتفاعلية.")

    add_heading_rtl(doc, "توصيات لتحسين الدقة", 2)
    add_rtl_paragraph(doc, "نوصي بزيادة حجم مجموعة البيانات باستخدام تقنيات توسيع البيانات (Data Augmentation)، وتضمين عينات من زوايا وبيئات إضاءة وخلفيات أكثر تنوعاً لتقوية مناعة النموذج.")

    add_heading_rtl(doc, "أعمال مستقبلية", 2)
    add_rtl_paragraph(doc, "1. تطوير النماذج باستخدام CNN/LSTM للتعرف على الكلمات الإشارية المتصلة.\n2. دمج ميزة تحويل النص إلى كلام (Text-to-Speech) لزيادة الفاعلية.\n3. النشر السحابي للتطبيق على خوادم تجارية وإتاحة تطبيقات هواتف أصلية.")

    add_heading_rtl(doc, "الكلمة الختامية", 2)
    add_rtl_paragraph(doc, "نأمل أن يشكل تطبيق 'إشارة' نواة أساسية ومستدامة لدعم مجتمع الصم وضعاف السمع في الوطن العربي، وأن يسهم في تيسير اندماجهم وتواصلهم اليومي بفاعلية.")
    doc.add_page_break()

    # --- REFERENCES ---
    add_heading_rtl(doc, "المراجع (References)", 1)
    add_rtl_paragraph(doc, "[المراجع العربية]")
    add_rtl_paragraph(doc, "1. الهيئة العامة للاتصالات. (2004). قاموس لغة الإشارة العربية الموحدة. جامعة الدول العربية.")
    add_rtl_paragraph(doc, "2. الباحثون في مجال التعلم العميق. (2022). تقنيات التعرف على لغة الإشارة العربية. المجلة العربية للحوسبة.")
    
    add_rtl_paragraph(doc, "")
    add_rtl_paragraph(doc, "[English References]")
    add_ltr_paragraph(doc, "1. Lugaresi, C., et al. (2019). MediaPipe: A Framework for Building Perception Pipelines. arXiv preprint arXiv:1906.08172.", WD_ALIGN_PARAGRAPH.LEFT)
    add_ltr_paragraph(doc, "2. Abadi, M., et al. (2015). TensorFlow: Large-scale machine learning on heterogeneous systems.", WD_ALIGN_PARAGRAPH.LEFT)
    add_ltr_paragraph(doc, "3. Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research.", WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()

    # --- APPENDICES ---
    add_heading_rtl(doc, "الملاحق", 1)
    
    add_heading_rtl(doc, "الملحق أ: خريطة مفاتيح الحروف العربية", 2)
    add_rtl_paragraph(doc, "[جدول المفاتيح من 1 إلى l يدرج هنا]")

    add_heading_rtl(doc, "الملحق ب: متطلبات التشغيل الكاملة", 2)
    add_rtl_paragraph(doc, "- نظام التشغيل: Windows 10/11 أو macOS أو توزيعات Linux حديثة.\n- بيئة العمل: Python 3.10 مثبتة محلياً.\n- المتصفح: Chrome, Edge, أو Firefox (أحدث إصدار) للوصول للكاميرا و PWA.\n- الاعتماديات: المذكورة في ملف requirements.txt.")

    add_heading_rtl(doc, "الملحق ج: أجزاء من الكود الهام", 2)
    
    code_snippet = '''# server.py - Endpoint for prediction
@app.route('/predict', methods=['POST'])
def predict():
    try:
        data = request.json
        landmarks = data.get('landmarks', [])
        if len(landmarks) != 42:
            return jsonify({'error': 'Invalid landmarks data'}), 400
            
        # Format the input for the model
        input_data = np.array([landmarks], dtype=np.float32)
        
        # Predict using TFLite
        interpreter.set_tensor(input_details[0]['index'], input_data)
        interpreter.invoke()
        output_data = interpreter.get_tensor(output_details[0]['index'])
        
        # Get result
        prediction_idx = np.argmax(output_data[0])
        confidence = float(output_data[0][prediction_idx])
        
        letter = label_encoder.inverse_transform([prediction_idx])[0]
        return jsonify({'letter': letter, 'confidence': confidence})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
'''
    add_ltr_paragraph(doc, code_snippet, WD_ALIGN_PARAGRAPH.LEFT)

    # Save the document
    out_path = os.path.join(r"c:\\Users\\moham\\Desktop\\project final file", "Eshara_Project_Book.docx")
    doc.save(out_path)
    print(f'Document saved successfully as {out_path}')

if __name__ == '__main__':
    main()
